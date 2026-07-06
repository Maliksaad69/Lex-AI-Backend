from pypdf import PdfReader
from docx import Document
import shutil
from typing import Tuple, List
import textract
import fitz
import re
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional

from pathlib import Path
# ---- Tesseract OCR (optional) ----
try:
    import pytesseract
    from PIL import Image
    import io

    # Point to Tesseract executable on Windows
    _tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _tp in _tesseract_paths:
        if Path(_tp).is_file():
            pytesseract.pytesseract.tesseract_cmd = _tp
            break

    def _ocr_page(page: fitz.Page, dpi: int = 300) -> str:
        """OCR a single page by rendering it to an image."""
        pix = page.get_pixmap(dpi=dpi)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img)

except ImportError:
    pytesseract = None

    def _ocr_page(page: fitz.Page, dpi: int = 300) -> str:
        return ""


UPLOAD_DIR = Path(__file__).resolve().parent.parent.parent / "uploads"

# Create the uploads directory if it doesn't exist
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def save_file(file) -> Path:
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    return file_path


MIN_TEXT_THRESHOLD = 30  # characters — below this, treat as scanned/image-only


def _clean_text(text: str) -> str:
    """Normalise whitespace in extracted text."""
    if not text:
        return ""
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _try_ocr(page: fitz.Page) -> Optional[str]:
    """Run OCR if Tesseract is available. Returns None if unavailable."""
    if pytesseract is None:
        return None
    try:
        return pytesseract.image_to_string(
            Image.open(io.BytesIO(page.get_pixmap(dpi=300).tobytes("png")))
        )
    except Exception:
        return None


def extract_text_from_pdf(file_path: str | Path) -> tuple[str, int]:
    """
    Extract text from a PDF.

    For pages with selectable text (most digital PDFs), uses direct text
    extraction — fast and accurate. For image-only / scanned pages, falls
    back to Tesseract OCR when available.

    Returns:
        full_text  (str)
        page_count (int)
    """
    file_path = Path(file_path)
    doc = fitz.open(str(file_path))

    pages: list[str] = []

    for page in doc:
        # 1. Try direct text extraction first (works for born-digital PDFs)
        text = _clean_text(page.get_text("text"))

        # 2. If barely any text, this page is likely scanned/image-only — try OCR
        if len(text) < MIN_TEXT_THRESHOLD:
            ocr_result = _try_ocr(page)
            if ocr_result:
                text = _clean_text(ocr_result)

        pages.append(text)

    doc.close()

    full_text = "\n\n".join(pages)
    return full_text, len(pages)


def extract_text_from_docx(file_path: str | Path) -> str:
    file_path = Path(file_path)
    doc = Document(file_path)

    content = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]
            content.append(" | ".join(cells))

        content.append("")  # Blank line after each table

    return "\n".join(content)